import sqlite3  # SQLite kütüphanesi
import requests
from bs4 import BeautifulSoup
import streamlit as st
import json
import csv
from docx import Document
import io
import time
import threading


# Steam desteklenen diller
LANGUAGES = {
    "Tüm Diller": "all",
    "İngilizce": "english",
    "Türkçe": "turkish",
    "Almanca": "german",
    "Fransızca": "french",
    "İspanyolca": "spanish"
}

# Oyunun ismini çekme
def get_game_name(app_id):
    url = f"https://store.steampowered.com/api/appdetails?appids={app_id}"
    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        if data.get(str(app_id), {}).get('success', False):
            return data[str(app_id)]['data']['name']
    return "Böyle bir oyun yok..."


# İncelemeleri çekme
def get_steam_reviews(app_id,language,num_pages=10):
    base_url = f"https://steamcommunity.com/app/{app_id}/reviews/"
    reviews = []

    for page in range(1, num_pages + 1):
        url = f"{base_url}?p={page}&browsefilter=mostrecent&filterLanguage={language}"
        response = requests.get(url)
        if response.status_code != 200:
            st.error(f"Sayfa {page} alınamadı. Durum kodu: {response.status_code}")
            break
        soup = BeautifulSoup(response.text, 'html.parser')
        review_blocks = soup.find_all('div', class_='apphub_CardTextContent')
        for block in review_blocks:
            review_text = block.get_text(strip=True)
            reviews.append(review_text)

    return reviews


# Word olarak kaydetme
def save_as_word(app_id,language):
    game_name = get_game_name(app_id)
    reviews = get_steam_reviews(app_id,language)

    if not reviews:
        st.warning("Kaydedilecek bir inceleme bulunamadı.")
        return

    document = Document()
    if language=="all":
        document.add_heading(f"Reviews for '{game_name}' (App ID: {app_id})", level=1)
    else:
        language=str(language).capitalize()
        document.add_heading(f"{language} Reviews for '{game_name}' (App ID: {app_id})", level=1)

    for i, review in enumerate(reviews, start=1):
        document.add_paragraph(f"Review {i}:")
        document.add_paragraph(review)
        document.add_paragraph()

    byte_io = io.BytesIO()
    document.save(byte_io)
    byte_io.seek(0)

    st.success(f"📄 **{game_name}** oyununun incelemeleri Word formatında indirilebilir!")
    st.download_button(
        label="Word dosyasını indir",
        data=byte_io,
        file_name=f"{game_name}_reviews.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# JSON olarak kaydetme
def save_as_json(app_id,language):
    game_name = get_game_name(app_id)
    reviews = get_steam_reviews(app_id,language)

    if not reviews:
        st.warning("Kaydedilecek bir inceleme yok.")
        return

    language=str(language).capitalize()
    json_data = json.dumps({"App ID": app_id, "Game Name": game_name, "Reviews": reviews,"Language":language}, indent=4, ensure_ascii=False)

    st.success(f"📝 **{game_name}** oyununun incelemeleri JSON formatında indirilebilir!")
    st.download_button(
        label="JSON dosyasını indir",
        data=json_data,
        file_name=f"{game_name}_reviews.json",
        mime="application/json"
    )


import csv
import io


# CSV olarak kaydetme
def save_as_csv(app_id,language):
    if not app_id.isdigit():
        st.error("App ID bir sayı olmalı!")
        return

    game_name = get_game_name(app_id)
    reviews = get_steam_reviews(app_id, language)

    if not reviews:
        st.warning("Kaydedilecek bir inceleme yok.")
        return

    # CSV'yi bellekte oluştur (BytesIO kullanıyoruz!)
    byte_io = io.BytesIO()
    text_io = io.TextIOWrapper(byte_io, encoding="utf-8", newline="")

    # CSV yazma işlemi
    csv_writer = csv.writer(text_io)
    csv_writer.writerow(["App ID", "Game Name", "Review","Language"])

    for review in reviews:
        csv_writer.writerow([app_id, game_name, review, language])

    # Streamlit için dosyayı sıfırdan okumaya hazır hale getir
    text_io.flush()
    byte_io.seek(0)

    st.success(f"📊 **{game_name}** oyununun incelemeleri CSV formatında indirilebilir!")

    # Kullanıcıya indirme butonu sun
    st.download_button(
        label="CSV dosyasını indir",
        data=byte_io.getvalue(),
        file_name=f"{game_name}_reviews.csv",
        mime="text/csv"
    )

def get_game_image(app_id):
    return f"https://cdn.cloudflare.steamstatic.com/steam/apps/{app_id}/header.jpg"

def update_game_name():
    app_id = st.session_state.app_id_input.strip()  # Boşlukları temizle
    if app_id.isdigit() and app_id:
        st.session_state.game_name = get_game_name(app_id)
    else:
        st.session_state.game_name = "Geçerli bir App ID girin"

# Kullanıcı Arayüzü
st.title("Steam Review Scraper")

st.write("Oyunun App ID'sini giriniz: ")
st.text_input("App ID", key="app_id_input", on_change=update_game_name)

IDapp = st.session_state.app_id_input  # App ID'yi tekrar kullanmak için değişkene ata

language_s = st.selectbox("İncelemeleri hangi dilde almak istersiniz?", list(LANGUAGES.keys()),key="language_select")

selected_language = LANGUAGES[language_s]

if IDapp.isdigit():
    image_url = get_game_image(IDapp)
    st.image(image_url,use_container_width=True)

if "game_name" in st.session_state:
    st.markdown(f"🎮 **Oyun Adı:** `{st.session_state.game_name}`")
else:
    st.warning("App ID'yi yazın ve geçerli bir ID girildiğinden emin olun.")

# Butonları ortalamak için CSS
st.markdown(
    """
    <style>
    .centered {
        display: flex;
        justify-content: center;
        gap: 20px;
        margin-top: 20px;
    }
    .loading-icon {
        display: flex;
        justify-content: center;
        margin-top: 20px;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Butonları içine koyacağımız div başlıyor
st.markdown('<div class="centered">', unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)

with col1:
    if st.button("Word olarak kaydet"):
        with st.spinner("📄 Word dosyası oluşturuluyor..."):
            save_as_word(IDapp,selected_language)

with col2:
    if st.button("JSON olarak kaydet"):
        with st.spinner("📝 JSON dosyası oluşturuluyor..."):
            save_as_json(IDapp,selected_language)

with col3:
    if st.button("CSV olarak kaydet"):
        with st.spinner("📊 CSV dosyası oluşturuluyor..."):
            save_as_csv(IDapp,selected_language)

st.markdown('</div>', unsafe_allow_html=True)

st.write("Aradığınız Oyunun App ID'sini öğrenmek için tıklayın.")
st.markdown("[SteamDB](https://steamdb.info/apps/)", unsafe_allow_html=True)